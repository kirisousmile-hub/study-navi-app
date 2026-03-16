############################################
# 0) IMPORTS / ENV
############################################
import os
from dotenv import load_dotenv

# LangChain telemetry OFF
os.environ["ANONYMIZED_TELEMETRY"] = "1"
os.environ["CHROMA_TELEMETRY"] = "0"


############################################
# 1) Standard Library
############################################
import shutil
import hashlib
import traceback
import json
import re
import uuid
import ast
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Tuple


############################################
# 2) Third Party Libraries
############################################
import streamlit as st

st.set_page_config(
    page_title="Study Navi",
    layout="wide"
)
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document as DocxDocument


############################################
# 3) LangChain / LLM
############################################
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader,
)

from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_core.messages import (
    SystemMessage,
    HumanMessage,
    AIMessage,
)

from langchain_core.documents import Document


############################################
# 4) Local Modules (core)
############################################

# knowledge map
from core.knowledge_map import (
    find_root_weakness,
    show_knowledge_map,
)

# utils
from core.utils import (
    infer_lesson_from_path,
    format_source_page,
    unique_by_source_page,
    count_turns,
)

# review system
from core.review import (
    load_review_cards,
    save_review_cards,
    compute_next_review_date,
    make_review_card,
)

# learning system
from core.learning import (
    load_learning_profile,
    generate_self_test,
    grade_answer,
    generate_weak_question,
    generate_today_mission,
    generate_ai_curriculum,
    explain_weakness,
    get_learning_level,
    get_weak_topics_sorted,
    update_learning_profile,
    add_learning_log,
    register_weak_point,
    load_weak_points,
)

# coach
from core.coach import (
    load_wall_memory,
    save_wall_memory,
    add_wall_fact,
    add_wall_summary,
    build_memory_block,
    summarize_wall_history,
    retrieve_hits,
    coach_reply,
    delete_wall_fact,
    delete_wall_summary,
)

# vector database
from core.vector_db import (
    file_fingerprint,
    generate_chunk_id,
    get_main_dir,
    get_registry_dir,
    load_registry,
    save_registry,
    is_file_indexed,
    mark_file_indexed,
    get_db,
    has_main_index,
    build_or_update_vectorstore,
)

# rag
from core.rag import prepare_documents

############################################
# ENV / MODEL CONFIG
############################################

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
MODEL_NAME = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

############################################
# GLOBAL LLM（1つだけ生成）
############################################

LLM = ChatOpenAI(
    model=MODEL_NAME,
    temperature=0.2,
    api_key=OPENAI_API_KEY
)

LLM_CREATIVE = ChatOpenAI(
    model=MODEL_NAME,
    temperature=0.4,
    api_key=OPENAI_API_KEY
)
LEARNING_DEPS = {
    "llm": LLM,
    "load_learning_profile": load_learning_profile,
    "update_learning_profile": update_learning_profile,
    "add_learning_log": add_learning_log,
    "register_weak_point": register_weak_point,
}

############################################
# GLOBAL EMBEDDINGS
############################################

EMBEDDINGS = OpenAIEmbeddings(
    model="text-embedding-3-small",
    api_key=OPENAI_API_KEY
)
db = get_db(EMBEDDINGS)

############################################
# 1) CONSTANTS / KEYS
############################################
# # なぜ：Streamlitは上から順に実行。キー/定数は「どこでも参照できる」ように最上段で固定
APP_TITLE = "Study Navi"

# ディレクトリ
PERSIST_DIR = "vectorstore_main"       # # なぜ：Chromaの本体DB保存先（run_idで分割）
REGISTRY_DIR = "vectorstore_registry"  # # なぜ：重複登録を防ぐための「登録済み台帳」保存先（run_idで分割）
TMP_UPLOAD_DIR = "tmp_uploads"         # # なぜ：アップロードファイルを一時保存してLoaderに渡すため

# ローカル教材フォルダ（自動取り込み）
LECTURES_PDF_DIR = Path("data/lectures_pdf")  # # なぜ：教材PDF（textbook扱い）
NOTES_DIR = Path("data/notes")                # # なぜ：自作メモ（notes扱い）

# session_state keys
RUN_ID_KEY = "run_id"              # # なぜ：DB領域をrun単位で分けて「完全初期化」を簡単にする
WALL_KEY = "wall_history"          # # なぜ：壁打ちチャット履歴
WALL_SUMMARY_KEY = "wall_summary"  # # なぜ：壁打ちまとめ
WALL_HITS_KEY = "wall_hits"        # # なぜ：壁打ちの検索ヒット（根拠候補）

# 壁打ち履歴の最大ターン
TURN_LIMIT = 30                    # # なぜ：LLMに渡す履歴が長すぎるとコスト/遅延/脱線が増えるため
# 自動要約を走らせるターン数
AUTO_SUMMARY_TURN = 12


############################################
# 2) UTILITIES (small)
############################################
def ensure_dirs():
    """# なぜ：起動時に必要な保存先が無いと失敗するので、先に作る"""
    Path(TMP_UPLOAD_DIR).mkdir(parents=True, exist_ok=True)
    Path(get_main_dir()).mkdir(parents=True, exist_ok=True)
    Path(get_registry_dir()).mkdir(parents=True, exist_ok=True)

def clear_tmp_uploads():
    """# なぜ：アップロードの差し替え時に古いファイルが残ると混乱するため"""
    if Path(TMP_UPLOAD_DIR).exists():
        shutil.rmtree(TMP_UPLOAD_DIR, ignore_errors=True)
    Path(TMP_UPLOAD_DIR).mkdir(parents=True, exist_ok=True)


def rerank_docs(question: str, docs: List[Document], embeddings, top_k: int = 4) -> List[Document]:
    if not docs:
        return []

    q_emb = embeddings.embed_query(question)
    texts = [d.page_content[:800] for d in docs]
    doc_embs = embeddings.embed_documents(texts)

    scores = cosine_similarity([q_emb], doc_embs)[0]
    scored = list(zip(scores, docs))
    scored.sort(key=lambda x: x[0], reverse=True)

    return [d for _, d in scored[:top_k]]


def retrieve_hits(
    query: str,
    db,
    embeddings,
    k: int = 5,
    only_textbook: bool = True
) -> List[Document]:
    search_kwargs = {"k": int(k) * 3}

    if only_textbook:
        search_kwargs["filter"] = {"category": "textbook"}

    retriever = db.as_retriever(search_kwargs=search_kwargs)
    raw_hits = retriever.invoke(query)

    reranked = rerank_docs(query, raw_hits, embeddings, top_k=int(k))

    return unique_by_source_page(reranked, int(k))

def format_sources(docs: List[Document]) -> str:
    """# なぜ：sources表示の重複を除いて読みやすくする"""
    seen = set()
    lines = []
    for d in docs:
        label = format_source_page(d.metadata)
        if label in seen:
            continue
        seen.add(label)
        lines.append(f"- {label}")
    return "\n".join(lines) if lines else "- (なし)"

def save_uploaded_files(uploaded_files) -> List[Path]:
    """# なぜ：StreamlitのUploadedFileはそのままLoaderに渡せないので一旦保存する"""
    ensure_dirs()
    clear_tmp_uploads()
    saved_paths: List[Path] = []
    for uf in uploaded_files:
        p = Path(TMP_UPLOAD_DIR) / uf.name
        with open(p, "wb") as f:
            f.write(uf.getbuffer())
        saved_paths.append(p)
    return saved_paths

def collect_local_files() -> List[Path]:
    """# なぜ：教材/メモをフォルダに置くだけで自動投入できるようにする"""
    paths: List[Path] = []
    if LECTURES_PDF_DIR.exists():
        paths.extend(sorted(LECTURES_PDF_DIR.glob("*.pdf")))
    if NOTES_DIR.exists():
        paths.extend(sorted(NOTES_DIR.glob("*.txt")))
        paths.extend(sorted(NOTES_DIR.glob("*.md")))
    return paths

############################################
# 3) LOAD / SPLIT / IDS
############################################
def load_docx(path: Path) -> List[Document]:
    """# なぜ：LangChainのdocx loaderは環境差で落ちることがあるため自前で読む"""
    doc = DocxDocument(str(path))
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text, metadata={"source": path.name, "path": str(path), "page": 1})]

def load_one_file(path: Path) -> List[Document]:
    """# なぜ：各形式の読み込み＋metadata付与を1箇所に集約して事故を減らす"""
    ext = path.suffix.lower()

    lesson = infer_lesson_from_path(path)
    is_lecture_pdf = (LECTURES_PDF_DIR in path.parents) and (ext == ".pdf")

    base_meta = {
        "path": str(path),
        "lesson": lesson,
        "category": "textbook" if is_lecture_pdf else "notes",
    }

    if ext == ".pdf":
        loader = PyPDFLoader(str(path))
        docs = loader.load()
        for d in docs:
            d.metadata["source"] = f"lectures/{path.name}"
            d.metadata.update(base_meta)
        return docs

    if ext == ".txt":
        for enc in ("utf-8", "utf-8-sig", "cp932"):
            try:
                loader = TextLoader(str(path), encoding=enc)
                docs = loader.load()
                for d in docs:
                    d.metadata["source"] = f"notes/{path.name}"
                    d.metadata.update(base_meta)
                return docs
            except Exception:
                continue
        raise ValueError(f"TXTの読み込みに失敗しました: {path.name}")

    # ※必要ならここに .csv .docx を追加
    raise ValueError(f"未対応形式: {ext}")

def split_docs(docs: List[Document]) -> List[Document]:
    """# なぜ：PDFをそのまま入れると長すぎて検索精度が落ちるのでchunk化する"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=150,
        separators=["\n\n", "\n", "。", " ", ""],
    )
    chunks = splitter.split_documents(docs)
    for c in chunks:
        c.metadata["source"] = c.metadata.get("source", "unknown")
        if "page" not in c.metadata:
            c.metadata["page"] = None
    return chunks




############################################
# 4) DB / REGISTRY
############################################
def get_run_id() -> str:
    """# なぜ：DB領域をrun単位で分け、完全初期化を簡単にする"""
    rid = st.session_state.get(RUN_ID_KEY)
    if rid:
        return rid
    rid = hashlib.md5(os.urandom(16)).hexdigest()[:8]
    st.session_state[RUN_ID_KEY] = rid
    return rid








############################################
# 5) RAG (answer_with_rag)
############################################
def answer_with_rag(
    question: str,
    k: int = 4,
    only_textbook: bool = False,
    lesson_filter: str | None = None
) -> Tuple[str, List[Document]]:
    """# なぜ：『質問→根拠→次の一手』を最小構成で安定させる"""
    db = get_db(EMBEDDINGS)

    search_kwargs = {"k": int(k) * 3}

    where = None

    if only_textbook and lesson_filter:
        where = {
            "$and": [
                {"category": "textbook"},
                {"lesson": lesson_filter}
            ]
        }

    elif only_textbook:
        where = {"category": "textbook"}

    elif lesson_filter:
        where = {"lesson": lesson_filter}

    if where:
        search_kwargs["filter"] = where

    retriever = db.as_retriever(search_kwargs=search_kwargs)
    raw_hits = retriever.invoke(question)

    # rerank追加
    reranked = rerank_docs(question, raw_hits, top_k=int(k))

    hits = unique_by_source_page(reranked, int(k))

    context = "\n\n".join(
        [f"[{i}] {format_source_page(d.metadata)}\n{d.page_content}" for i, d in enumerate(hits, start=1)]
    )


    prompt = f"""あなたは「学習ナビ」です。
以下の「参照コンテキスト」だけに基づいて回答してください。
推測で断定しない。分からなければ「不明」と言い、確認手順を提案する。

# ユーザーの質問
{question}

# 参照コンテキスト
{context}

# 出力フォーマット（必ず守る）
【結論】
- （1〜3行）

【根拠（参照した資料の要点）】
- （必ず番号[1]などを交えて）
- （可能なら必ず「PDF名 p.X」を含める）

【次の一手（最短3つ）】
1.
2.
3.
"""
    return LLM.invoke(prompt).content, hits

############################################
# LEARNING PROFILE
############################################

def show_weak_heatmap():

    profile = load_learning_profile()

    if not profile:
        st.info("弱点データがまだありません")
        return

    st.subheader("🔥 弱点ヒートマップ")

    weak_topics = get_weak_topics_sorted(profile)

    for i, (_, topic) in enumerate(weak_topics[:5]):

        st.write(f"{i+1}位 : {topic}")

def show_learning_dashboard():

    profile = load_learning_profile()

    if not profile:
        st.info("まだ理解度データがありません")
        return

    st.subheader("📊 学習ダッシュボード")

    topics = []
    scores = []

    for topic, data in profile.items():

        total = data["total"]
        correct = data["correct"]

        if total == 0:
            score = 0
        else:
            score = int((correct / total) * 100)

        topics.append(topic[:20])
        scores.append(score)

    chart_data = {
        "topic": topics,
        "score": scores
    }

    st.bar_chart(chart_data, x="topic", y="score")

def generate_adaptive_question():

    profile = load_learning_profile()

    if not profile:
        return "まだ弱点データがありません"

    weak_topics = get_weak_topics_sorted(profile)
    topic = weak_topics[0][1]

    # RAG検索
    hits = retrieve_hits(topic, db, EMBEDDINGS, k=3)

    context = "\n\n".join(
        [d.page_content[:600] for d in hits]
    )

    prompt = f"""
Pythonの次のトピックについて理解確認問題を1つ作ってください。

トピック
{topic}

教材
{context}

条件
・短い問題
・コード理解問題
"""

    return LLM_CREATIVE.invoke(prompt).content

def generate_next_question():

    db = get_db(EMBEDDINGS)
    profile = load_learning_profile()

    if not profile:
        return None

    weak_topics = get_weak_topics_sorted(profile)
    topic = weak_topics[0][1]

    hits = retrieve_hits(topic, db, EMBEDDINGS, k=3)

    context = "\n".join(
        [d.page_content[:500] for d in hits]
    )

    prompt = f"""
        次の条件で、Pythonの理解確認問題を1つ作ってください。

        # トピック
        {topic}

        # 教材
        {context}

        # 条件
        ・短い問題
        ・コード理解問題
        ・出力は必ずJSONのみ
        ・説明文やコードフェンスは不要

        # 出力形式
        {{
        "question": "問題文",
        "reference": "模範回答",
        "explanation": "解説"
        }}
        """

    result = LLM_CREATIVE.invoke(prompt).content.strip()
    result = result.replace("```json", "").replace("```", "").strip()

    start = result.find("{")
    end = result.rfind("}")
    if start != -1 and end != -1 and start < end:
        result = result[start:end + 1]

    data = json.loads(result)
    data["topic"] = topic

    return data


############################################
# 8) UI
############################################

st.title(APP_TITLE)
st.caption("アップロードしたあなたのメモ/提出物を元に、根拠つきで回答し、次の一手を3つ提案します。")

if not OPENAI_API_KEY:
    st.warning("OPENAI_API_KEY が設定されていません。.env を確認してください。")

# 共通初期化
ensure_dirs()

if WALL_KEY not in st.session_state:
    st.session_state[WALL_KEY] = []

if WALL_SUMMARY_KEY not in st.session_state:
    st.session_state[WALL_SUMMARY_KEY] = None

if WALL_HITS_KEY not in st.session_state:
    st.session_state[WALL_HITS_KEY] = []

# =========================
# Sidebar：メンテ＆データ投入
# =========================
with st.sidebar:
    st.subheader("🔧 メンテ")

    # ---- 現在のDB状況を表示（押しすぎ防止）----
    try:
        db_count = get_db(EMBEDDINGS)._collection.count()
    except Exception:
        db_count = 0

    st.caption(f"📦 現在のインデックス: {db_count} チャンク")

    # registry登録済みファイル数
    try:
        reg = load_registry()
        st.caption(f"🧾 登録済みファイル: {len(reg)} 件")
    except Exception:
        st.caption("🧾 登録済みファイル: -")

    if st.session_state.get("just_reset", False):
        st.success("完全初期化しました。PDFを1つ入れて再検証してください。")
        st.session_state["just_reset"] = False

    if st.button("🧨 完全初期化（DB + registry + tmp）"):
        shutil.rmtree(get_main_dir(), ignore_errors=True)
        shutil.rmtree(get_registry_dir(), ignore_errors=True)
        shutil.rmtree(TMP_UPLOAD_DIR, ignore_errors=True)
        Path(TMP_UPLOAD_DIR).mkdir(parents=True, exist_ok=True)

        for key_name in ["retriever", "db"]:
            st.session_state.pop(key_name, None)

        st.success("DB/registry/tmp を完全初期化しました。")
        st.rerun()

    st.divider()
    st.header("① データ投入")

    uploaded = st.file_uploader(
        "PDF/DOCX/TXT/CSVをアップロード",
        type=["pdf", "docx", "txt", "csv"],
        accept_multiple_files=True
    )

    col_a, col_b = st.columns(2)
    with col_a:
        k = st.number_input("検索件数 k", min_value=2, max_value=10, value=4, step=1)
    with col_b:
        st.write("")

############################################
# TABS UI
############################################
tabs = st.tabs([
    "📚 学習",
    "🧠 壁打ち",
    "📝 復習",
    "📊 分析",
    "📂 教材管理"
])

tab_learn, tab_coach, tab_review, tab_analytics, tab_material = tabs





with tab_learn:
    st.header("📚 学習")

    # =========================
    # ② 質問する（RAG）
    # =========================
    st.divider()
    st.header("② 質問する")

    only_textbook = st.checkbox("教材（lectures_pdf）だけ検索する", value=True)
    lesson_filter = st.text_input(
        "Lessonフィルター（例: 13）",
        value=""
    )
    question = st.text_input("質問（例：Chroma永続化のしくみを自分の言葉で説明したい）", value="")

    ask_disabled = (not question.strip()) or (not has_main_index(db))

    if st.button("質問する", disabled=ask_disabled):
        try:
            with st.spinner("検索＆回答中..."):
                ans, hits = answer_with_rag(
                question,
                k=int(k),
                only_textbook=only_textbook,
                lesson_filter=lesson_filter
            )

            st.session_state["last_question"] = question
            st.session_state["last_answer"] = ans
            st.session_state["last_hits"] = hits
        except Exception as e:
            st.error(f"回答に失敗: {e}")
            st.code(traceback.format_exc())

    last_q = st.session_state.get("last_question")
    last_answer = st.session_state.get("last_answer")
    last_hits = st.session_state.get("last_hits")

    if last_answer:
        left, right = st.columns([2, 1])
        with left:
            st.subheader("回答")
            st.write(last_answer)

            if st.button("📝（直前の回答）この質問をカード化"):
                try:
                    cards = load_review_cards()
                    card = make_review_card(topic=last_q, answer=last_answer, hits=last_hits)
                    cards.append(card)
                    save_review_cards(cards)
                    st.success(f"カード化しました！ id={card['id']}")
                    st.rerun()
                except Exception as e:
                    st.error(f"カード化に失敗: {e}")
                    st.code(traceback.format_exc())

            # ↓↓↓ここに追加↓↓↓
            if st.button("🧠 この内容で自己テスト作成"):
                try:
                    test = generate_self_test(last_q, last_hits, LLM_CREATIVE)
                    st.session_state["self_test"] = test
                    st.rerun()
                except Exception as e:
                    st.error(f"自己テスト作成に失敗しました: {e}")

            # テスト表示
            test = st.session_state.get("self_test")

            if test:
                st.subheader("自己テスト")

                for i, q in enumerate(test["questions"]):

                    st.markdown(f"### Q{i+1}. {q}")

                    user_key = f"user_answer_{i}"

                    user_answer = st.text_area(
                        "あなたの回答",
                        key=user_key
                    )

                    if st.button(f"採点する Q{i+1}", key=f"grade_{i}"):

                        result = grade_answer(
                            q,   # topic
                            q,   # question
                            user_answer,
                            test["answers"][i],
                            LEARNING_DEPS["llm"],
                            LEARNING_DEPS["update_learning_profile"],
                            LEARNING_DEPS["add_learning_log"],
                            LEARNING_DEPS["register_weak_point"]
                        )
                        st.markdown("### AI採点")
                        st.write(result)

        with right:
            st.subheader("参照（sources）")
            st.markdown(format_sources(last_hits))

    # 壁打ち履歴の見える化（RAGの下に置くのが分かりやすい）
    turns, msgs = count_turns(st.session_state[WALL_KEY])
    st.caption(f"壁打ち履歴: {turns} / {TURN_LIMIT} ターン（メッセージ {msgs}件）")

    st.divider()
    st.subheader("AIトレーニング")

    if st.button("弱点トレーニング開始"):

        q = generate_weak_question(load_weak_points, LLM_CREATIVE)

        if q:
            st.session_state["duo_question"] = q
        else:
            st.warning("まだ弱点がありません")

    duo_q = st.session_state.get("duo_question")

    if duo_q:

        st.subheader("AIトレーニング問題")

        st.write(duo_q)

        duo_answer = st.text_area("あなたの回答", key="duo_answer")

        if st.button("回答を採点", key="duo_grade_btn"):

            topic = duo_q

            result = grade_answer(
                duo_q,   # topic
                duo_q,   # question
                duo_answer,
                "Pythonの正しい説明",
                LLM,
                update_learning_profile,
                add_learning_log,
                register_weak_point
            )
            st.write(result)



    turns, msgs = count_turns(st.session_state[WALL_KEY])
    st.caption(f"壁打ち履歴: {turns} / {TURN_LIMIT} ターン（メッセージ {msgs}件）")

    st.divider()



    # 今日のミッション
    st.subheader("🎯 今日のミッション")

    if st.button("ミッション生成"):

        mission = generate_today_mission(
            LEARNING_DEPS["load_learning_profile"],
            LEARNING_DEPS["llm"]
        )

        st.session_state["mission"] = mission

    if "mission" in st.session_state:
        st.write(st.session_state["mission"])

    st.divider()

    # AIカリキュラム
    st.subheader("🤖 AIカリキュラム")
    # =========================
    # 次に学ぶトピック
    # =========================

    st.subheader("🧭 次に学ぶトピック")

    def recommend_next_topic():

        profile = load_learning_profile()

        if not profile:
            return "まだ学習データがありません"

        weakest, prereq = find_root_weakness(profile)

        prompt = f"""
    Python学習コーチとして答えてください。

    現在の弱点
    {weakest}

    前提知識
    {prereq}

    次に学ぶべきトピックを
    1つだけ提案してください。

    短く答えてください。
    """

        return LLM.invoke(prompt).content


    if st.button("次に学ぶ内容を提案"):

        topic = recommend_next_topic()

        st.success(topic)

    if st.button("今日の学習メニュー生成"):

        curriculum = generate_ai_curriculum(load_learning_profile, LLM)

        st.session_state["ai_curriculum"] = curriculum

        

    curriculum = st.session_state.get("ai_curriculum")

    if curriculum:
        st.write(curriculum)

    st.subheader("🚀 AI学習モード")

    if st.button("AI学習開始"):

        q = generate_next_question()
        st.session_state["loop_question"] = q

    loop_data = st.session_state.get("loop_question")

    if loop_data:

        st.write("### 問題")
        st.write(loop_data["question"])

        loop_answer = st.text_area("あなたの回答", key="loop_answer")

        if st.button("回答を採点", key="loop_grade_btn"):

            result = grade_answer(
                loop_data["topic"],
                loop_data["question"],
                loop_answer,
                loop_data["reference"],
                LLM,
                update_learning_profile,
                add_learning_log,
                register_weak_point
            )

            st.write("### AI採点")
            st.write(result)

            st.markdown("### 模範回答")
            st.write(loop_data["reference"])

            st.markdown("### 解説")
            st.write(loop_data["explanation"])

        if st.button("次の問題へ", key="loop_next_btn"):

            st.session_state["loop_question"] = generate_next_question()
            st.rerun()

    # =========================
    # AIドリルモード
    # =========================

    st.subheader("🧠 AIドリル")

    if "drill_question" not in st.session_state:
        st.session_state["drill_question"] = None


    def generate_drill_question(topic):

        prompt = f"""
    Python学習者向けに
    {topic} の理解度を確認する問題を1つ作ってください。

    条件
    ・短い問題
    ・初心者向け
    ・答えは書かない
    """

        return LLM_CREATIVE.invoke(prompt).content


    if st.button("問題を出す"):

        profile = load_learning_profile()

        if profile:
            weak_topics = get_weak_topics_sorted(profile)
            topic = weak_topics[0][1]
        else:
            topic = "基礎プログラミング"

        st.session_state["drill_question"] = generate_drill_question(topic)


    if st.session_state["drill_question"]:

        st.write("### 問題")
        st.write(st.session_state["drill_question"])

        answer = st.text_area("あなたの回答")

        if st.button("回答を採点", key="drill_grade_btn"):

            grading_prompt = f"""
    次の問題の回答を採点してください。

    問題
    {st.session_state["drill_question"]}

    回答
    {answer}

    出力
    ・正解かどうか
    ・改善ポイント
    """

            result = LLM.invoke(grading_prompt).content

            st.write("### AI採点")
            st.write(result)

            if st.button("次の問題", key="drill_next_btn"):

                st.session_state["drill_question"] = generate_drill_question(topic)
                st.rerun()

    # =========================
    # 自動次問題
    # =========================

    profile = load_learning_profile()

    if profile:
        weak_topics = get_weak_topics_sorted(profile)
        next_topic = weak_topics[0][1]
    else:
        next_topic = "基礎プログラミング"


    st.info(f"次の問題（弱点トピック）：{next_topic}")

    if st.button("次の問題へ"):
        st.session_state["loop_question"] = generate_next_question()
        st.rerun()


with tab_coach:
    st.header("🧠 壁打ち")

    colL, colR = st.columns([2, 1])

    with colR:
        wall_mode = st.selectbox(
            "学習フェーズ",
            ["A: 用語理解", "B: 設計理解", "C: コード理解"],
            index=0
        )
        wall_only_textbook = st.checkbox(
            "教材だけ参照",
            value=True,
            key="wall_only_textbook"
        )
        wall_k = st.number_input(
            "壁打ち検索k",
            min_value=2,
            max_value=8,
            value=4,
            step=1,
            key="wall_k"
        )
        use_long_memory = st.checkbox(
            "長期記憶を使う",
            value=False,
            key="use_long_memory"
        )

        st.divider()
        st.subheader("🧠 覚えさせるメモ（永続）")
        st.caption("保存済みメモ（直近）")
        st.caption("保存先の整理")
        st.caption("・壁打ち履歴：session_state")
        st.caption("・長期記憶：wall_memory.json")
        st.caption("・復習カード：review_cards.json")
        st.text_area(
            "saved_memory",
            build_memory_block(limit=10, include_facts=True, include_summaries=True),
            height=180,
            disabled=True
        )

        st.divider()
        st.subheader("🗂 保存済みメモを個別削除")

        mem_data = load_wall_memory()

        facts = mem_data.get("facts", [])
        summaries = mem_data.get("summaries", [])

        if facts:
            st.caption("固定メモ")
            fact_options = {
                f"{f['id']} | {f['text'][:40]}": f["id"]
                for f in facts
            }
            selected_fact_label = st.selectbox(
                "削除する固定メモ",
                options=list(fact_options.keys()),
                key="delete_fact_select"
            )

            if st.button("固定メモを削除", key="delete_fact_btn"):
                ok = delete_wall_fact(fact_options[selected_fact_label])
                if ok:
                    st.success("固定メモを削除しました")
                    st.rerun()
                else:
                    st.warning("削除対象が見つかりませんでした")

        if summaries:
            st.caption("学習要約")
            summary_options = {
                f"{s['id']} | {s['text'][:40].replace(chr(10), ' ')}": s["id"]
                for s in summaries
            }
            selected_summary_label = st.selectbox(
                "削除する学習要約",
                options=list(summary_options.keys()),
                key="delete_summary_select"
            )

            if st.button("学習要約を削除", key="delete_summary_btn"):
                ok = delete_wall_summary(summary_options[selected_summary_label])
                if ok:
                    st.success("学習要約を削除しました")
                    st.rerun()
                else:
                    st.warning("削除対象が見つかりませんでした")

        mem_text = st.text_input(
            "覚えてほしいこと（例：合言葉はリンゴ）",
            key="mem_text"
        )

        colm1, colm2 = st.columns(2)
        with colm1:
            if st.button("➕ メモを保存", key="save_mem_btn"):
                if mem_text.strip():
                    f = add_wall_fact(mem_text)
                    st.success(f"保存しました id={f['id']}")
                    st.rerun()
                else:
                    st.warning("空です")

        with colm2:
            if st.button("🗑 メモ全消し（危険）", key="clear_mem_btn"):
                save_wall_memory({"facts": [], "summaries": []})
                st.warning("全メモを削除しました")
                st.rerun()

    with colL:
        st.caption("あなたが喋る → 根拠を差し込む → 質問で掘る、の順で進めます。")

        chat_area = st.container(height=600, border=True)

        with chat_area:
            for m in st.session_state[WALL_KEY]:
                with st.chat_message(m["role"]):
                    st.write(m["content"])

        user_msg = st.chat_input(
            "いま何を復習したい？（例：for文 / 関数 / 例外 / import / 合言葉確認）"
        )

        if user_msg:
            st.session_state[WALL_KEY].append(
                {"role": "user", "content": user_msg}
            )

            hits = retrieve_hits(
                user_msg,
                db,
                EMBEDDINGS,
                k=int(wall_k),
                only_textbook=wall_only_textbook
            )
            st.session_state[WALL_HITS_KEY] = hits

            assistant_msg = coach_reply(
                st.session_state[WALL_KEY],
                hits,
                wall_mode,
                LLM,
                use_long_memory=use_long_memory
            )

            st.session_state[WALL_KEY].append(
                {"role": "assistant", "content": assistant_msg}
            )

            turns = len(st.session_state[WALL_KEY]) // 2

            if turns >= AUTO_SUMMARY_TURN:
                summary_text = summarize_wall_history(
                    st.session_state[WALL_KEY],
                    hits,
                    LLM
                )

                add_wall_summary(summary_text)

                st.session_state[WALL_SUMMARY_KEY] = summary_text

                st.session_state[WALL_KEY] = [
                    {
                        "role": "assistant",
                        "content": "ここまでの壁打ち内容を自動要約して長期記憶に保存しました。続きをどうぞ。"
                    }
                ]

            st.rerun()

        hits = st.session_state.get(WALL_HITS_KEY, [])
        if hits:
            with st.expander("参照（sources）"):
                st.markdown(format_sources(hits))

    st.divider()
    colA, colB = st.columns(2)

    with colA:
        if st.button("🧾 この壁打ちをまとめる"):
            hist = st.session_state[WALL_KEY]
            hits = st.session_state.get(WALL_HITS_KEY, [])

            summary_text = summarize_wall_history(
                hist,
                hits,
                LLM
            )

            st.session_state[WALL_SUMMARY_KEY] = summary_text
            add_wall_summary(summary_text)

            st.rerun()

    with colB:
        if st.button("🗑 壁打ちをリセット"):
            st.session_state[WALL_KEY] = []
            st.session_state.pop(WALL_SUMMARY_KEY, None)
            st.session_state.pop(WALL_HITS_KEY, None)
            st.rerun()

    summary = st.session_state.get(WALL_SUMMARY_KEY)
    if summary:
        st.subheader("まとめ")
        st.write(summary)
        if st.button("📝 この壁打ちまとめをカード化"):
            cards = load_review_cards()
            card = make_review_card(
                topic="壁打ちまとめ",
                answer=summary,
                hits=st.session_state.get(WALL_HITS_KEY, []),
            )
            cards.append(card)
            save_review_cards(cards)
            st.success(f"カード化しました！ id={card['id']}")
            st.rerun()

    st.caption(f"壁打ち履歴: {len(st.session_state[WALL_KEY])} メッセージ")

with tab_review:
    st.header("📝 復習")
    # =========================
    # ③ 復習する
    # =========================

    cards = load_review_cards()
    if not cards:
        st.info("まだカードがありません。まずは②で質問 → 回答 → 『📝この質問をカード化』を押してください。")
    else:
        today = datetime.now().strftime("%Y-%m-%d")
        only_due = st.checkbox("今日の復習だけ表示する（next_review_date <= 今日）", value=True)

        def is_due(card: dict) -> bool:
            d = card.get("next_review_date")
            return (d is None) or (d <= today)

        filtered = [c for c in cards if (is_due(c) if only_due else True)]
        st.write(f"カード数: {len(cards)} / 表示: {len(filtered)}（今日={today}）")

        sort_key = st.selectbox("並び順", ["next_review_dateが古い順", "作成日時が新しい順"], index=0)
        if sort_key == "next_review_dateが古い順":
            filtered.sort(key=lambda c: c.get("next_review_date") or "0000-00-00")
        else:
            filtered.sort(key=lambda c: c.get("created_at") or "", reverse=True)

        options = [
            f"{c.get('id','????')} | {c.get('next_review_date','-')} | {c.get('topic','(no topic)')[:50]}"
            for c in filtered
        ]
        sel = st.selectbox("カードを選択", options, index=0)
        sel_id = sel.split("|")[0].strip()
        card = next((c for c in cards if c.get("id") == sel_id), None)

        if not card:
            st.error("カードが見つかりませんでした（データ不整合の可能性）。")
        else:
            st.subheader("問題（自分で答えてから開く）")
            st.write(card.get("question") or card.get("topic"))

            show_answer = st.checkbox("答えを表示する", value=False)
            if show_answer:
                st.subheader("答え（保存された回答）")
                st.write(card.get("answer", ""))

                st.subheader("参照（sources）")
                srcs = card.get("sources", [])
                if srcs:
                    st.markdown("\n".join([f"- {s}" for s in srcs]))
                else:
                    st.write("(なし)")

            st.divider()
            st.subheader("採点して次回復習日を更新")

            col1, col2, col3 = st.columns(3)

            def update_score(new_score: int):
                card["score"] = new_score
                card["last_review_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                card["next_review_date"] = compute_next_review_date(new_score)

                for i, c in enumerate(cards):
                    if c.get("id") == card.get("id"):
                        cards[i] = card
                        break

                save_review_cards(cards)
                st.success(f"更新しました：score={new_score} / next={card['next_review_date']}")
                st.rerun()

            with col1:
                if st.button("0 😵 無理（明日）"):
                    update_score(0)
            with col2:
                if st.button("1 🤔 微妙（2日後）"):
                    update_score(1)
            with col3:
                if st.button("2 ✅ できた（1週間後）"):
                    update_score(2)

            st.divider()
            st.subheader("管理")
            if st.button("このカードを削除（危険）"):
                cards2 = [c for c in cards if c.get("id") != card.get("id")]
                save_review_cards(cards2)
                st.warning("削除しました。")
                st.rerun()

            st.divider()
            st.subheader("弱点トピック")

            weak = load_weak_points()

            if weak:
                weak_sorted = sorted(weak, key=lambda x: x["count"], reverse=True)

                for w in weak_sorted[:5]:
                    st.write(f"{w['topic']} （{w['count']}回ミス）")

            else:
                st.caption("まだ弱点はありません")

with tab_analytics:
    st.header("📊 分析")

    show_learning_dashboard()

    show_weak_heatmap()

    show_knowledge_map()

    st.subheader("🧠 弱点分析")

    if st.button("弱点の原因を分析"):

        explanation = explain_weakness(
            LEARNING_DEPS["load_learning_profile"],
            LEARNING_DEPS["llm"]
        )

        st.session_state["weak_explain"] = explanation

    explanation = st.session_state.get("weak_explain")

    if explanation:
        st.write(explanation)

    st.divider()

    st.subheader("🏆 学習レベル")

    level = get_learning_level(load_learning_profile)

    st.write(level)

    st.subheader("🧑‍🏫 AI学習アドバイス")

    if st.button("学習状況を分析"):

        profile = load_learning_profile()

        if not profile:
            st.info("まだ学習データがありません")
        else:

            summary = []

            for topic, data in profile.items():

                total = data["total"]
                correct = data["correct"]

                score = correct / total if total else 0

                summary.append(f"{topic}:{round(score*100)}%")

            prompt = f"""
Python学習コーチとして
次の学習データを分析してください。

{summary}

・現在の理解度
・弱点
・次の学習アドバイス

を短く説明してください。
"""

            advice = LLM.invoke(prompt).content

            st.write(advice)

with tab_material:
    st.header("📂 教材管理")
    # =========================
    # ① インデックス作成
    # =========================
    local_paths = collect_local_files()

    if st.button("インデックス作成（追加）"):
        try:
            saved_paths: List[Path] = []
            if uploaded:
                saved_paths = save_uploaded_files(uploaded)

            local_paths = collect_local_files()

            candidate_paths: List[Path] = []
            candidate_paths.extend(saved_paths)
            candidate_paths.extend(local_paths)

            target_paths: List[Path] = []
            skipped: List[str] = []

            for p in candidate_paths:
                fp = file_fingerprint(p)
                if is_file_indexed(fp):
                    skipped.append(p.name)
                    continue
                target_paths.append(p)

            raw_docs: List[Document] = []
            for p in target_paths:
                raw_docs.extend(load_one_file(p))

            chunks = split_docs(raw_docs)

            if len(chunks) == 0:
                st.success(
                    f"追加完了：アップロード {len(saved_paths)}ファイル / "
                    f"ローカル {len(local_paths)}ファイル / "
                    f"対象 {len(target_paths)}ファイル / "
                    f"スキップ {len(skipped)}ファイル / "
                    f"0チャンク"
                )
                if skipped:
                    st.info("スキップ: " + ", ".join(skipped[:20]) + ("..." if len(skipped) > 20 else ""))
            else:
                before = get_db(EMBEDDINGS)._collection.count()

                build_or_update_vectorstore(chunks)

                after = get_db(EMBEDDINGS)._collection.count()
                st.info(f"Chroma count: {before} -> {after} (+{after - before})")

                for p in target_paths:
                    mark_file_indexed(file_fingerprint(p), p)

                st.success(
                    f"追加完了：アップロード {len(saved_paths)}ファイル / "
                    f"ローカル {len(local_paths)}ファイル / "
                    f"対象 {len(target_paths)}ファイル / "
                    f"スキップ {len(skipped)}ファイル / "
                    f"{len(chunks)}チャンク"
                )

                if skipped:
                    st.info("スキップ: " + ", ".join(skipped[:20]) + ("..." if len(skipped) > 20 else ""))

        except Exception as e:
            st.error(f"失敗: {e}")
            st.code(traceback.format_exc())






