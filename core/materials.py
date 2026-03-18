import shutil
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    CSVLoader,
    PyPDFLoader,
    TextLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter

from core.utils import infer_lesson_from_path


LECTURES_PDF_DIR = Path("data/lectures_pdf")
NOTES_DIR = Path("data/notes")
TMP_UPLOAD_DIR = "tmp_uploads"

TEXT_ENCODINGS = ("utf-8", "utf-8-sig", "cp932")


def make_base_meta(path: Path) -> dict:
    lesson = infer_lesson_from_path(path)
    is_lecture_pdf = (LECTURES_PDF_DIR in path.parents) and (path.suffix.lower() == ".pdf")

    return {
        "path": str(path),
        "lesson": lesson,
        "category": "textbook" if is_lecture_pdf else "notes",
    }


def make_source_label(path: Path) -> str:
    ext = path.suffix.lower()
    is_lecture_pdf = (LECTURES_PDF_DIR in path.parents) and (ext == ".pdf")

    if is_lecture_pdf:
        return f"lectures/{path.name}"

    if TMP_UPLOAD_DIR in str(path):
        return f"uploads/{path.name}"

    return f"notes/{path.name}"


def apply_common_metadata(docs: List[Document], path: Path, base_meta: dict) -> List[Document]:
    source = make_source_label(path)

    for doc in docs:
        doc.metadata["source"] = source
        doc.metadata.update(base_meta)

    return docs


def load_docx(path: Path) -> List[Document]:
    doc = DocxDocument(str(path))
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "path": str(path),
                "page": 1,
            },
        )
    ]


def save_uploaded_files(uploaded_files) -> List[Path]:
    Path(TMP_UPLOAD_DIR).mkdir(parents=True, exist_ok=True)

    saved_paths: List[Path] = []
    for uploaded_file in uploaded_files:
        path = Path(TMP_UPLOAD_DIR) / uploaded_file.name
        with open(path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        saved_paths.append(path)

    return saved_paths


def clear_tmp_uploads(tmp_upload_dir: str = TMP_UPLOAD_DIR) -> None:
    """古い一時アップロードを削除して、空の一時フォルダを作り直す"""
    tmp_path = Path(tmp_upload_dir)

    if tmp_path.exists():
        shutil.rmtree(tmp_path, ignore_errors=True)

    tmp_path.mkdir(parents=True, exist_ok=True)


def collect_local_files() -> List[Path]:
    paths: List[Path] = []

    if LECTURES_PDF_DIR.exists():
        paths.extend(sorted(LECTURES_PDF_DIR.glob("*.pdf")))

    if NOTES_DIR.exists():
        paths.extend(sorted(NOTES_DIR.glob("*.txt")))
        paths.extend(sorted(NOTES_DIR.glob("*.md")))
        paths.extend(sorted(NOTES_DIR.glob("*.docx")))
        paths.extend(sorted(NOTES_DIR.glob("*.csv")))

    return paths


def load_text_like_file(path: Path) -> List[Document]:
    for encoding in TEXT_ENCODINGS:
        try:
            loader = TextLoader(str(path), encoding=encoding)
            return loader.load()
        except UnicodeDecodeError:
            continue
        except Exception:
            continue

    raise ValueError(f"{path.suffix.upper()} の読み込みに失敗しました: {path.name}")


def load_csv_with_fallback(path: Path) -> List[Document]:
    for encoding in TEXT_ENCODINGS:
        try:
            loader = CSVLoader(str(path), encoding=encoding)
            return loader.load()
        except UnicodeDecodeError:
            continue
        except Exception:
            continue

    raise ValueError(f"CSVの読み込みに失敗しました: {path.name}")


def load_one_file(path: Path) -> List[Document]:
    ext = path.suffix.lower()
    base_meta = make_base_meta(path)

    if ext == ".pdf":
        docs = PyPDFLoader(str(path)).load()
        return apply_common_metadata(docs, path, base_meta)

    if ext in (".txt", ".md"):
        docs = load_text_like_file(path)
        return apply_common_metadata(docs, path, base_meta)

    if ext == ".docx":
        docs = load_docx(path)
        return apply_common_metadata(docs, path, base_meta)

    if ext == ".csv":
        docs = load_csv_with_fallback(path)
        return apply_common_metadata(docs, path, base_meta)

    raise ValueError(f"未対応形式: {ext}")


def split_docs(docs: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=150,
        separators=["\n\n", "\n", "。", " ", ""],
    )

    chunks = splitter.split_documents(docs)

    for chunk in chunks:
        chunk.metadata["source"] = chunk.metadata.get("source", "unknown")
        if "page" not in chunk.metadata:
            chunk.metadata["page"] = None

    return chunks